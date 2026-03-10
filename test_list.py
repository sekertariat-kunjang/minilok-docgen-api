from docxtpl import DocxTemplate, RichText
from docx import Document

# 1. Create a template programmatically (or we can just load an empty one)
doc = Document()
doc.add_paragraph('Test Single Placeholder List:')
p = doc.add_paragraph('{{ test_list }}', style='List Number')
doc.save('test_list_template.docx')

# 2. Render with docxtpl
tpl = DocxTemplate('test_list_template.docx')

# Try standard string with newlines
print("Rendering standard string with newlines...")
tpl.render({'test_list': "Item A\nItem B\nItem C"})
tpl.save('test_list_output1.docx')

# Try RichText
print("Rendering RichText with newlines...")
rt = RichText()
rt.add("Item X\nItem Y\nItem Z")
tpl.render({'test_list': rt})
tpl.save('test_list_output2.docx')
