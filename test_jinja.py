from docxtpl import DocxTemplate
from docx import Document

# Create template manually or programmatically
doc = Document()
doc.add_paragraph('Jinja List Test:')
p = doc.add_paragraph('{%p for item in ai_prosedur_list %}')
p2 = doc.add_paragraph('{{ item }}', style='List Number')
p3 = doc.add_paragraph('{%p endfor %}')
doc.save('test_jinja_template.docx')

# Render
tpl = DocxTemplate('test_jinja_template.docx')

# Simulate what main.py will do: 
# The UI sends a string with newlines. The backend splits it into a list.
ui_input = "Cuci tangan dengan sabun\nBilas dengan air mengalir\nKeringkan dengan handuk"
ai_prosedur_list = ui_input.split('\n')

tpl.render({'ai_prosedur_list': ai_prosedur_list})
tpl.save('test_jinja_output.docx')
print("Successfully rendered Jinja List")
