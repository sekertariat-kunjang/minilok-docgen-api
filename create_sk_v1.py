import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_sk_template():
    doc = docx.Document()
    
    # Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("PEMERINTAH KABUPATEN KEDIRI\nDINAS KESEHATAN\nUPTD PUSKESMAS MINILOK")
    run.bold = True
    run.font.size = Pt(14)
    
    doc.add_paragraph("Alamat: Jl. Raya No. 123, Kediri").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("-" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("KEPUTUSAN KEPALA UPTD PUSKESMAS MINILOK")
    run.bold = True
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("NOMOR: {{nomor_sk}}").bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("TENTANG\n").bold = True
    p.add_run("{{ai_tentang}}").bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("KEPALA UPTD PUSKESMAS MINILOK,").bold = True
    
    # Menimbang
    p = doc.add_paragraph()
    p.add_run("Menimbang : ").bold = True
    p.add_run("a. bahwa {{ai_menimbang}};\n")
    p.add_run("            b. bahwa berdasarkan pertimbangan sebagaimana dimaksud dalam huruf a, perlu menetapkan Keputusan Kepala UPTD Puskesmas Minilok tentang {{ai_tentang}};")
    
    # Mengingat
    p = doc.add_paragraph()
    p.add_run("Mengingat : ").bold = True
    p.add_run("1. {{ai_mengingat}};\n")
    p.add_run("            2. Peraturan Menteri Kesehatan Nomor 43 Tahun 2019 tentang Pusat Kesehatan Masyarakat;")
    
    # Menetapkan
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("MEMUTUSKAN").bold = True
    
    p = doc.add_paragraph()
    p.add_run("MENETAPKAN : ").bold = True
    p.add_run("KEPUTUSAN KEPALA UPTD PUSKESMAS MINILOK TENTANG {{ai_tentang}}.")
    
    p = doc.add_paragraph()
    p.add_run("KESATU : ").bold = True
    p.add_run("{{ai_menetapkan}}")
    
    p = doc.add_paragraph()
    p.add_run("KEDUA : ").bold = True
    p.add_run("Keputusan ini berlaku sejak tanggal ditetapkan dengan ketentuan apabila di kemudian hari terdapat kekeliruan akan diadakan perbaikan sebagaimana mestinya.")
    
    # Signature
    doc.add_paragraph("\n\n")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Ditetapkan di Kediri\n")
    p.add_run("pada tanggal: {{tanggal}}\n")
    p.add_run("KEPALA UPTD PUSKESMAS MINILOK,\n\n\n\n")
    p.add_run("dr. BUDI SANTOSO").bold = True
    
    doc.save("templates/sk_template_v1.docx")

if __name__ == "__main__":
    create_sk_template()
