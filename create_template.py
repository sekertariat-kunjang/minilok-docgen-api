from docxtpl import DocxTemplate
import os

def create_complex_template():
    templates_dir = r"c:\Users\star\3D Objects\minlok\docgen-service\templates"
    os.makedirs(templates_dir, exist_ok=True)
    
    # We can't easily create a .docx with complex tables via docxtpl from scratch without a base file,
    # but we can use python-docx to build the "mold" and then save it.
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Header Table
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'LOGO'
    hdr_cells[1].text = '{{ nama_puskesmas }}\n{{ alamat_puskesmas }}'
    
    doc.add_paragraph()
    
    # Title
    title = doc.add_paragraph('{{ judul }}')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(14)
    
    # Info Table
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Table Grid'
    
    labels = ['No. Dokumen', 'No. Revisi', 'Tgl. Terbit', 'Halaman']
    vars = ['{{ nomor_dokumen }}', '{{ nomor_revisi }}', '{{ tanggal_terbit }}', '{{ halaman }}']
    
    for i in range(4):
        info_table.rows[i].cells[0].text = labels[i]
        info_table.rows[i].cells[1].text = vars[i]
        
    doc.add_paragraph()
    
    # Content Sections
    sections = [
        ('1. Pengertian', '{{ pengertian }}'),
        ('2. Tujuan', '{{ tujuan }}'),
        ('3. Kebijakan', '{{ kebijakan }}'),
        ('4. Referensi', '{{ referensi }}')
    ]
    
    for head, body in sections:
        p_head = doc.add_paragraph(head)
        p_head.runs[0].bold = True
        doc.add_paragraph(body)
        
    # Procedures Looping Table (The Complex Part)
    doc.add_paragraph('5. Prosedur / Langkah-langkah').runs[0].bold = True
    
    proc_table = doc.add_table(rows=2, cols=2)
    proc_table.style = 'Table Grid'
    
    # Docxtpl loop tags
    proc_table.rows[0].cells[0].text = 'No'
    proc_table.rows[0].cells[1].text = 'Langkah-langkah'
    
    # row 1 will be the loop row
    # Use j2 tags: {% tr for item in prosedur_list %} ... {% tr endfor %}
    proc_table.rows[1].cells[0].text = '{% tr for item in prosedur_list %}{{ item.num }}'
    proc_table.rows[1].cells[1].text = '{{ item.step }}{% tr endfor %}'
    
    doc.add_paragraph()
    doc.add_paragraph('6. Unit Terkait: {{ unit_terkait }}')
    
    # Signature
    doc.add_paragraph()
    sig = doc.add_paragraph('Disetujui Oleh,\nKepala Puskesmas\n\n\n\n{{ kepala_puskesmas }}\nNIP. {{ nip_kepala }}')
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    template_path = os.path.join(templates_dir, "complex_sop.docx")
    doc.save(template_path)
    print(f"Created complex template at: {template_path}")

if __name__ == "__main__":
    create_complex_template()
