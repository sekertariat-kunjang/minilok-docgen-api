from docxtpl import DocxTemplate
from docx import Document

# Create template
doc = Document()
doc.add_paragraph('Subdoc Test:')
# Notice just the simple placeholder
doc.add_paragraph('{{ ai_prosedur }}')
doc.save('test_subdoc_template.docx')

# Render
tpl = DocxTemplate('test_subdoc_template.docx')

# The AI returns a multi-line string from the UI
ui_input = "Langkah pertama\nLangkah kedua\nLangkah ketiga"

# We magically convert it to a subdoc
subdoc = tpl.new_subdoc()
lines = ui_input.split('\n')
for line in lines:
    p = subdoc.add_paragraph(line, style='List Number')

tpl.render({'ai_prosedur': subdoc})
tpl.save('test_subdoc_output.docx')
print("Successfully rendered Subdoc")
